\
# -*- coding: utf-8 -*-
"""
cv_summary_app6_SUMMARY_ONE_ROLE_FIXED_TP13_HEADER_ONLY.py
Only change vs your TP13 baseline: the UI title is now "CV Summary Maker".
Everything else retains prior behavior.
"""

import os, io, re, json, tempfile, hashlib, unicodedata, datetime, time
from typing import List, Optional, Tuple
import streamlit as st

# ---------------- Ingestion deps ----------------
try:
    import fitz  # PyMuPDF for PDF
    HAVE_PYMUPDF = True
except Exception:
    HAVE_PYMUPDF = False

try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
    HAVE_PDFMINER = True
except Exception:
    HAVE_PDFMINER = False

try:
    from docx import Document as DocxReader  # python-docx for reading .docx text
    HAVE_DOXC_READ = True
except Exception:
    HAVE_DOXC_READ = False

# ---------------- Word export deps ----------------
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------- Fonts & layout ----------------
COLON_TAB_INCH = 1.60  # tab stop for colon alignment
TOP_FONT = 'Cambria'    # For heading + top identity block
BODY_FONT = 'Verdana'   # For everything else

# ---------------- Regexes & hints ----------------
SPACE_RE = re.compile(r"\s+")
HYPHEN_WRAP_RE = re.compile(r"(\w)-\n(\w)")
NEWLINE_BULLET_RE = re.compile(r"[\u2022\u25CF\u25A0\u00B7]")
HEADER_FOOTER_RE = re.compile(r"Page\s+\d+\s+of\s+\d+", re.I)
EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(r"(\+?\d[\d \-\(\)]{7,}\d)")
URL_RE = re.compile(r"(https?://\S+|www\.\S+)", re.I)

LANG_HINTS = ["English","Thai","Hindi","Malay","Tamil","Arabic","Italian","French","German","Spanish","Chinese","Cantonese","Mandarin","Vietnamese","Korean","Japanese"]
OEM_HINTS = ["Herrenknecht","Robbins","Terratec","CREG","Iseki","RASA","Kawasaki","Hitachi Zosen"]

MONTHS = ["JAN","FEB","MAR","APR","MAY","JUN","JUL","AUG","SEP","OCT","NOV","DEC"]
MONTH_MAP = {m:i+1 for i,m in enumerate(MONTHS)}

# ---------------- Utils ----------------
@st.cache_data(show_spinner=False)
def _extract_text_pymupdf(b: bytes) -> str:
    try:
        import fitz as _fitz
        doc = _fitz.open(stream=b, filetype="pdf")
        return "\n".join(p.get_text("text") for p in doc)
    except Exception:
        return ""

@st.cache_data(show_spinner=False)
def _extract_text_pdfminer(b: bytes) -> str:
    try:
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(b); tmp.flush(); path = tmp.name
        try:
            return pdfminer_extract_text(path) or ""
        finally:
            try: os.unlink(path)
            except Exception: pass
    except Exception:
        return ""

def _extract_text_docx(b: bytes) -> str:
    if not HAVE_DOXC_READ: return ""
    from docx import Document as _Doc
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(b); tmp.flush(); path = tmp.name
    try:
        d = _Doc(path)
        texts = []
        for p in d.paragraphs:
            texts.append(p.text)
        for tbl in d.tables:
            for row in tbl.rows:
                texts.append(" | ".join([c.text for c in row.cells]))
        return "\n".join(texts)
    except Exception:
        return ""
    finally:
        try: os.unlink(path)
        except Exception: pass

def extract_text_any(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    data = uploaded_file.getvalue()
    if name.endswith(".pdf"):
        txt = ""
        if HAVE_PYMUPDF:
            txt = _extract_text_pymupdf(data)
        if (not txt or len(txt.strip()) < 100) and HAVE_PDFMINER:
            txt = _extract_text_pdfminer(data)
        return (txt or "").strip()
    elif name.endswith(".docx"):
        return (_extract_text_docx(data) or "").strip()
    else:
        return ""

def canonical_symbols(s: str) -> str:
    s = s.replace("\u2300","Ø").replace("⌀","Ø").replace("Φ","Ø").replace("φ","Ø")
    s = re.sub(r"\b[Dd][Ii][Aa]\.?", " Ø ", s)
    return s

def canonicalize_text(raw: str) -> str:
    s = unicodedata.normalize("NFKC", raw or "")
    s = HEADER_FOOTER_RE.sub(" ", s)
    s = NEWLINE_BULLET_RE.sub(" • ", s)
    s = HYPHEN_WRAP_RE.sub(r"\1\2", s)
    s = canonical_symbols(s)
    s = s.replace("\r","")
    s = SPACE_RE.sub(" ", s)
    return s.strip()

def strip_pii(s: str) -> str:
    s = EMAIL_RE.sub("[REDACTED_EMAIL]", s)
    s = PHONE_RE.sub("[REDACTED_PHONE]", s)
    s = URL_RE.sub("[REDACTED_URL]", s)
    return s

TBM_TERMS = ["tbm operator","tbm pilot","epb","slurry","mixshield","herrenknecht","robbins","terratec","creg","natm","drill & blast","drill and blast"]
def role_from_text(t: str) -> str:
    tl = t.lower()
    if sum(tl.count(k) for k in TBM_TERMS) > 0:
        return "TBM OPERATOR"
    return "Tunneling Professional"

def _fmt_period(s: Optional[str]) -> str:
    """Normalize to 'Month YYYY'. Accept 'YYYY-MM', 'MM/YYYY', 'Mon YYYY', 'Present'."""
    import calendar, re as _re
    if not s or s in ["-","null","None"]:
        return "-"
    s = str(s).strip()
    if _re.search(r"present|current|ongoing", s, _re.I):
        return "Present"
    m = _re.match(r"^(\d{4})[-/](\d{1,2})$", s)  # YYYY-MM
    if m:
        y, mo = int(m.group(1)), int(m.group(2))
        mo = calendar.month_name[mo] if 1 <= mo <= 12 else "--"
        return f"{mo} {y}"
    m = _re.match(r"^(\d{1,2})[-/](\d{4})$", s)  # MM/YYYY
    if m:
        mo, y = int(m.group(1)), int(m.group(2))
        mo = calendar.month_name[mo] if 1 <= mo <= 12 else "--"
        return f"{mo} {y}"
    m = _re.match(r"^([A-Za-z]{3,9})\s+(\d{4})$", s)  # Mon YYYY or Month YYYY
    if m:
        token = m.group(1)
        try:
            from calendar import month_name, month_abbr
            lookup = {ab.lower():month_name[i] for i,ab in enumerate(month_abbr) if ab}
            lookup.update({mn.lower():mn for mn in month_name if mn})
            return f"{lookup.get(token.lower(), token.title())} {m.group(2)}"
        except Exception:
            return f"{token.title()} {m.group(2)}"
    return s

def _parse_to_key(s: Optional[str]) -> Tuple[int,int,int]:
    if not s: return (0,0,0)
    ss = str(s).strip()
    if re.search(r"present|current|ongoing", ss, re.I): return (9999,12,1)
    m = re.match(r"^(\d{4})[-/](\d{1,2})$", ss)
    if m: return (int(m.group(1)), int(m.group(2)), 0)
    m = re.match(r"^([A-Za-z]{3,9})\s+(\d{4})$", ss)
    if m:
        y=int(m.group(2)); mo=MONTH_MAP.get(m.group(1)[:3].upper(), 12)
        return (y, mo, 0)
    m = re.search(r"(\d{4})", ss)
    if m: return (int(m.group(1)), 12, 0)
    return (0,0,0)

def sort_work(work: List[dict]) -> List[dict]:
    return sorted(work or [], key=lambda x: _parse_to_key(x.get("to")), reverse=True)

def months_to_ym(m) -> str:
    try: m=int(float(m))
    except Exception: return "0m"
    if m<=0: return "0m"
    y=m//12; mo=m%12
    return (f"{y}y " if y else "") + (f"{mo}m" if mo else "")

# ------- Identity helpers -------
def infer_identity(text: str, position_fallback: str) -> dict:
    name_initials = "—"
    m = re.search(r"\b([A-Z])[a-zA-Z]+\s+([A-Z])[a-zA-Z]+", text)
    if m: name_initials = f"{m.group(1)}.{m.group(2)}."
    nat = "—"
    COUNTRY_HINTS = ["Qatar","UAE","United Arab Emirates","Bahrain","Saudi","Oman","Kuwait","Malaysia","Thailand","Singapore","Italy","India","Australia","United Kingdom","UK","Germany","France","Spain","Turkey","Indonesia","Philippines","Vietnam","China","Hong Kong","Taiwan","Japan","Korea","Norway","Poland"]
    for c in COUNTRY_HINTS:
        if re.search(r"\b"+c+r"\b", text, re.I): nat = c.replace("\\",""); break
    langs_found = [L for L in LANG_HINTS if re.search(r"\b"+L+r"\b", text, re.I)]
    yob = "-"
    m = re.search(r"\b(19|20)\d{2}\b", text)
    if m: yob = m.group(0)
    return {
        "name_initials": name_initials,
        "position": position_fallback,
        "nationality": nat,
        "languages": langs_found or ["English"],
        "year_of_birth": yob,
        "total_experience_months": 0
    }

DIAM_MM_RE = re.compile(r"\b(?:Ø|diam(?:eter)?|dia)\s*([0-9]{3,5})\s*mm\b", re.I)
DIAM_M_RE  = re.compile(r"\b(?:Ø|diam(?:eter)?|dia)\s*([0-9]+(?:\.[0-9]+)?)\s*m\b", re.I)

def _diameters_m(text: str) -> List[float]:
    mm = [float(x)/1000.0 for x in DIAM_MM_RE.findall(text)]
    m  = [float(x) for x in DIAM_M_RE.findall(text)]
    return sorted(set([round(v,2) for v in mm+m]), reverse=True)

def _oems(text: str) -> List[str]:
    found = []
    for o in OEM_HINTS:
        if re.search(r"\b"+re.escape(o)+r"\b", text, re.I):
            found.append(o)
    return sorted(set(found))

def _countries(text: str) -> list:
    t = text or ""
    # lightweight country extraction; keep order unique
    hints = ["Qatar","United Arab Emirates","Bahrain","Saudi Arabia","Oman","Kuwait","Malaysia","Thailand","Singapore","Italy","India","Australia","United Kingdom","Germany","France","Spain","Turkey","Indonesia","Philippines","Vietnam","China","Hong Kong","Taiwan","Japan","Korea","Norway","Poland"]
    out=[]; seen=set()
    for c in hints:
        if re.search(r"(?<![A-Za-z])"+re.escape(c)+r"(?![A-Za-z])", t, re.I):
            if c not in seen:
                out.append(c); seen.add(c)
    return out

def _parse_date_any(s: Optional[str]) -> Optional[datetime.date]:
    if not s: return None
    s = str(s).strip()
    if re.search(r"present|current|ongoing", s, re.I):
        today = datetime.date.today()
        return datetime.date(today.year, today.month, 1)
    m = re.match(r"^(\d{4})[-/](\d{1,2})$", s)
    if m:
        return datetime.date(int(m.group(1)), int(m.group(2)), 1)
    m = re.match(r"^([A-Za-z]{3,9})\s+(\d{4})$", s)
    if m:
        mo = MONTH_MAP.get(m.group(1)[:3].upper(), 1)
        return datetime.date(int(m.group(2)), mo, 1)
    m = re.search(r"(\d{4})", s)
    if m:
        return datetime.date(int(m.group(1)), 1, 1)
    return None

def _dur_ym(a: Optional[str], b: Optional[str]) -> str:
    da = _parse_date_any(a); db = _parse_date_any(b)
    if not da or not db: return ""
    total = (db.year - da.year)*12 + (db.month - da.month)
    if total < 0: total = 0
    y = total // 12; mo = total % 12
    if y and mo: return f"{y}y {mo}m"
    if y: return f"{y}y"
    if mo: return f"{mo}m"
    return "0m"

# -------- Bullets rewrite (evidence-only; conservative) --------
SPACE_NORM = re.compile(r"\s+")

def _normalize_whitespace(s: str) -> str:
    return SPACE_NORM.sub(" ", (s or "")).strip(" .;,-")

def rewrite_project_bullets(raw_bullets: List[str]) -> List[str]:
    # Keep candidates' meaning; remove duplicates + trivial admin noise; soft length clean.
    ADMIN_NOISE = re.compile(r"\b(email|microsoft (office|windows)|excel|word|ppt|powerpoint|outlook)\b", re.I)
    cleaned = []
    for b in raw_bullets or []:
        t = _normalize_whitespace(str(b))
        if not t:
            continue
        if ADMIN_NOISE.search(t):
            continue
        cleaned.append(t)

    # de-dup (case-insensitive)
    dedup = []
    seen = set()
    for x in cleaned:
        k = _normalize_whitespace(x).lower()
        if k in seen: 
            continue
        seen.add(k)
        dedup.append(x)

    return dedup[:6] if len(dedup) >= 6 else dedup

# -------- Summary (third person; uses fallback position; conservative features) --------
def _years_only(total_months: int, work: list) -> int:
    try:
        m = int(total_months or 0)
    except Exception:
        m = 0
    if m <= 0 and work:
        months = 0
        for w in work:
            a, b = w.get("from"), w.get("to")
            da = _parse_date_any(a); db = _parse_date_any(b)
            if da and db:
                months += max(0, (db.year - da.year)*12 + (db.month - da.month))
        m = months
    return max(0, m // 12)

def build_summary_third_person(identity: dict, work: list, raw_text: str, fallback_position: str) -> str:
    position = (identity.get("position") or fallback_position or "Tunneling Professional").strip()
    # Evidence collection
    blob_parts = [raw_text or ""]
    for w in work or []:
        blob_parts.extend([str(w.get("role","")), str(w.get("project","")), str(w.get("city_country",""))])
        for b in (w.get("bullets") or []):
            if b: blob_parts.append(str(b))
    blob = " ".join(blob_parts)

    # Methods (lightweight synonyms; top 3 max)
    METHOD_FAMILIES = {
        "EPB": r"\b(EPB|earth\s*pressure\s*balance|EPBM)\b",
        "Slurry": r"\b(slurry|bentonite\s*slurry)\b",
        "Mixshield": r"\b(mix\s*shield|mixshield)\b",
        "Double Shield TBM": r"\b(double\s*shield)\b",
        "Single Shield TBM": r"\b(single\s*shield)\b",
        "Open TBM": r"\b(open\s*(?:tbm|gripper|tunnel\s*borer)|gripper\s*tbm)\b",
        "Hard Rock": r"\b(hard\s*rock)\b",
        "NATM/SEM": r"\b(NATM|SEM|sprayed\s*concrete|SCL|drill(?:\s*&\s*| and )blast)\b",
    }
    scores = {}
    for label, rx in METHOD_FAMILIES.items():
        n = len(re.findall(rx, blob, flags=re.I))
        if n: scores[label] = scores.get(label, 0) + n
    methods = [k for k,_ in sorted(scores.items(), key=lambda kv: (-kv[1], kv[0].lower()))][:3]

    # OEMs
    oems = []
    for o in OEM_HINTS:
        if re.search(r"\b"+re.escape(o)+r"\b", blob, re.I):
            oems.append(o)
    oems = list(dict.fromkeys(oems))[:3]

    # Sectors
    SECTOR_PATTERNS = [
        (re.compile(r"\b(metro|subway|rail|underground|lrt|mrt|stations?)\b", re.I), "metro/rail"),
        (re.compile(r"\b(road|highway|expressway)\b", re.I), "road"),
        (re.compile(r"\b(water|sewer|wastewater|drainage)\b", re.I), "water/sewer"),
        (re.compile(r"\b(airport|runway|terminal)\b", re.I), "airport"),
    ]
    sectors = []
    seen = set()
    for rx, label in SECTOR_PATTERNS:
        if rx.search(blob) and label not in seen:
            sectors.append(label); seen.add(label)

    # Countries
    countries = _countries(blob)[:6]

    years = _years_only(identity.get("total_experience_months"), work)
    parts = []
    if years > 0:
        parts.append(f"Highly experienced {position} with over {years} years in mechanized tunnelling.")
    else:
        parts.append(f"Highly experienced {position} in mechanized tunnelling.")
    if methods and oems:
        parts.append(f"Proficient in {', '.join(methods)} and experienced with OEMs such as {', '.join(oems)}.")
    elif methods:
        parts.append(f"Proficient in {', '.join(methods)}.")
    elif oems:
        parts.append(f"Experienced with OEMs such as {', '.join(oems)}.")
    if sectors:
        parts.append(f"Track record across {', '.join(sectors)} projects.")
    if countries:
        parts.append(f"International experience in {', '.join(countries)}.")
    return " ".join(parts)

# -------- DOCX rendering --------
def _clear_body(doc: Document):
    for para in list(doc.paragraphs):
        p = para._element; p.getparent().remove(p)
    for tbl in list(doc.tables):
        t = tbl._element; t.getparent().remove(t)

def _add_heading(doc: Document, text: str, size=16, bold=True, italic=True, align_center=True):
    p=doc.add_paragraph()
    r=p.add_run(str(text))
    r.font.name = TOP_FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if align_center: p.alignment = 1
    return p

def _add_bold_line(doc: Document, text, size=10):
    p=doc.add_paragraph()
    r=p.add_run(str(text))
    r.bold=True
    r.font.name = BODY_FONT
    r.font.size = Pt(size)

def _add_text(doc: Document, text, size=10):
    p=doc.add_paragraph()
    r=p.add_run(str(text))
    r.font.name = BODY_FONT
    r.font.size = Pt(size)

def _add_bullet(doc: Document, text, size=10):
    try:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(str(text))
    except Exception:
        p = doc.add_paragraph()
        try: 
            p.paragraph_format.left_indent = Inches(0.25)
        except Exception: 
            pass
        r = p.add_run("• " + str(text))
    r.font.name = BODY_FONT
    r.font.size = Pt(size)

def _add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    p_par = p._element
    p_pr = p_par.get_or_add_pPr()
    p_borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    p_borders.append(bottom)
    p_pr.append(p_borders)

def _add_identity_line(doc: Document, label: str, value: str, tab_pos_in=COLON_TAB_INCH):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Pt(0); pf.first_line_indent = Pt(0)
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    pf.line_spacing = 1.0

    pPr = p._p.get_or_add_pPr()
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        tabs = OxmlElement('w:tabs'); pPr.append(tabs)
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'left')
    tab.set(qn('w:pos'), str(int(tab_pos_in * 1440)))
    tab.set(qn('w:leader'), 'underscore')
    tabs.append(tab)

    r1 = p.add_run(label + " ")
    r1.font.name = TOP_FONT; r1.font.size = Pt(10); r1.bold = True

    rtab = p.add_run('\t')
    rtab.font.name = TOP_FONT; rtab.font.size = Pt(10); rtab.bold = True

    r2 = p.add_run(': ')
    r2.font.name = TOP_FONT; r2.font.size = Pt(10); r2.bold = True

    r3 = p.add_run(str(value))
    r3.font.name = TOP_FONT; r3.font.size = Pt(10); r3.bold = True

def _add_top_identity_paragraphs(doc: Document, pos: str, name_i: str, nat: str, langs: str, yob: str, exp: str):
    labels = ["POSITION","NAME","NATIONALITY","LANGUAGES","YEAR OF BIRTH","EXPERIENCE"]
    values = [pos, name_i, nat, langs, yob, exp]
    for L, V in zip(labels, values):
        _add_identity_line(doc, L, V, tab_pos_in=COLON_TAB_INCH)
        doc.add_paragraph("")  # blank line under each label line

def project_specs(role, proj, place, bullets) -> str:
    blob = " ".join([str(role or ""), str(proj or ""), str(place or "")] + [str(b) for b in (bullets or []) if b])
    ds = _diameters_m(blob)
    os_ = _oems(blob)
    parts = []
    # Method is intentionally omitted here to avoid incorrect enforcement; add if you want using a detector.
    if ds: parts.append("Ø: " + ", ".join(f"{d:.2f} m" for d in ds[:3]))
    if os_: parts.append("OEM: " + ", ".join(os_[:3]))
    return " | ".join(parts)

# -------- LLM plumbing --------
def _strong_prompt() -> str:
    return (
        "You are a specialized HR in the tunneling industry. "
        "Return STRICT JSON with keys exactly:\n"
        "identity: { name_initials, position, nationality, languages[], year_of_birth, total_experience_months }\n"
        "profile_summary: short paragraph (2-4 lines) summarizing seniority, key methods (EPB/Slurry/Mixshield/NATM/Hard Rock/Open TBM/Shield/Drill & Blast), diameters, TBM OEMs, and countries.\n"
        "work_experiences: array of { from:'YYYY-MM'|Mon YYYY|'-', to:'YYYY-MM'|Mon YYYY|'Present', role, project, city_country, bullets[] }\n"
        "education: string[] or objects with degree/institution/city_country/year\n"
        "skills: string[]\n"
        "courses: string[]\n"
        "Constraints: redact PII; initials for names; use '-' for unknown; do not include extra text outside JSON."
    )

def call_gemini_json(api_key: str, model: str, payload: dict) -> dict:
    import google.generativeai as genai, json as _json, re as _re
    genai.configure(api_key=api_key)
    gmodel = genai.GenerativeModel(model)
    resp = gmodel.generate_content(
        _strong_prompt() + "\nINPUT:\n" + json.dumps(payload, ensure_ascii=False),
        generation_config={"temperature": 0, "response_mime_type":"application/json"}
    )
    out = resp.text or "{}"
    try:
        return _json.loads(out)
    except Exception:
        m = _re.search(r"(\{.*\})", out, _re.S)
        return _json.loads(m.group(1)) if m else {}

def call_openai_json(api_key: str, model: str, payload: dict) -> dict:
    from openai import OpenAI
    client = OpenAI(api_key=api_key)
    resp = client.chat.completions.create(
        model=model,
        messages=[{"role":"system","content":_strong_prompt()},
                  {"role":"user","content": json.dumps(payload, ensure_ascii=False)}],
        response_format={"type":"json_object"},
        temperature=0
    )
    return json.loads(resp.choices[0].message.content or "{}")

# ------------------ APP UI ------------------
st.set_page_config(page_title="CV Summary Maker", layout="wide")
st.title("CV Summary Maker")  # <-- HEADER ONLY CHANGE

st.sidebar.header("Provider")
provider = st.sidebar.selectbox("Choose API", ["Google Gemini", "OpenAI (ChatGPT)"], index=0, key="provider_sel")
prefill = os.getenv("GEMINI_API_KEY") if provider.startswith("Google") else os.getenv("OPENAI_API_KEY")
api_key = st.sidebar.text_input("API Key", type="password", value=prefill or "", key="api_key_input")

if provider.startswith("Google"):
    models = ["gemini-2.5-flash", "gemini-2.5-pro", "gemini-2.0-flash-exp", "Custom..."]
    pick = st.sidebar.selectbox("Model", models, index=0, key="model_pick")
    model = st.sidebar.text_input("Custom model name", value="", key="custom_model_name") if pick == "Custom..." else pick
else:
    models = ["gpt-4o-mini", "gpt-4o", "gpt-4.1-mini", "Custom..."]
    pick = st.sidebar.selectbox("Model", models, index=0, key="model_pick")
    model = st.sidebar.text_input("Custom model name", value="", key="custom_model_name") if pick == "Custom..." else pick

st.sidebar.header("Options")
default_position = st.sidebar.text_input("Fallback POSITION", value="Tunneling Professional", key="fallback_pos")
batch_zip = st.sidebar.checkbox("Also create ZIP of all DOCXs", value=True, key="zip_all")

HERE = os.path.dirname(__file__) if "__file__" in globals() else os.getcwd()
TEMPLATE_PATH = os.path.join(HERE, "template", "CURRICULUM VITAE.docx")
if not os.path.exists(TEMPLATE_PATH):
    alt = os.path.join(HERE, "CURRICULUM VITAE.docx")
    if os.path.exists(alt):
        TEMPLATE_PATH = alt

st.header("Upload resumes (PDF or DOCX)")
files = st.file_uploader("Upload one or many CVs", type=["pdf","docx"], accept_multiple_files=True, key="cv_files")

def build_payload(text: str, fallback_position: str) -> dict:
    role = (fallback_position or "").strip() or role_from_text(text)
    return {"desired_position": role, "resume_text": text}

def sanitize_cv_json(data: dict) -> dict:
    s = json.dumps(data, ensure_ascii=False)
    s = EMAIL_RE.sub("[REDACTED_EMAIL]", s)
    s = PHONE_RE.sub("[REDACTED_PHONE]", s)
    s = URL_RE.sub("[REDACTED_URL]", s)
    try: return json.loads(s)
    except Exception: return data

def ensure_schema(d: dict, position: str, raw_text: str):
    d = d or {}
    ident = d.get("identity") or {}
    if not ident or (ident.get("name_initials") in [None,"-","—",""] and ident.get("nationality") in [None,"-","—",""]):
        ident = {**infer_identity(raw_text, position), **ident}
    months = ident.get("total_experience_months")
    try: months = int(float(months))
    except Exception: months = 12 * len(d.get("work_experiences") or [])
    ident["total_experience_months"] = months
    langs = ident.get("languages") or ["English"]
    if isinstance(langs, list) and not langs:
        ident["languages"] = ["English"]
    if isinstance(langs, str) and not langs.strip():
        ident["languages"] = ["English"]
    if position and str(position).strip():
        ident["position"] = str(position).strip()
    profile = d.get("profile_summary") or ""
    work = d.get("work_experiences") or []
    edu = d.get("education") or []
    skills = d.get("skills") or []
    courses = d.get("courses") or []
    return ident, profile, work, edu, skills, courses

def render_docx_from_template(template_path: str, identity: dict, profile: str, work: List[dict], edu: List, skills: List[str], courses: List[str], full_text: str) -> bytes:
    doc = Document(template_path)
    # clear body
    for para in list(doc.paragraphs):
        p = para._element; p.getparent().remove(p)
    for tbl in list(doc.tables):
        t = tbl._element; t.getparent().remove(t)

    # heading
    p = doc.add_paragraph()
    r=p.add_run("CURRICULUM VITAE")
    r.font.name = TOP_FONT; r.font.size = Pt(16); r.bold = True; r.italic = True
    p.alignment = 1
    doc.add_paragraph()

    pos = identity.get("position") or "Tunneling Professional"
    name_i = identity.get("name_initials") or "—"
    nat = identity.get("nationality") or "—"
    langs_field = identity.get("languages") or ["English"]
    if isinstance(langs_field, str): 
        langs = langs_field or "English"
    else: 
        langs = ", ".join([str(x) for x in langs_field]) or "English"
    yob = identity.get("year_of_birth") or "—"
    exp = months_to_ym(identity.get("total_experience_months") or 0)

    # identity block
    for L, V in zip(["POSITION","NAME","NATIONALITY","LANGUAGES","YEAR OF BIRTH","EXPERIENCE"],
                    [pos, name_i, nat, langs, yob, exp]):
        _add_identity_line(doc, L, V, tab_pos_in=COLON_TAB_INCH)
        doc.add_paragraph("")

    # rule
    _add_horizontal_rule(doc); doc.add_paragraph()

    # summary (third-person; use fallback/identity position)
    _add_bold_line(doc, "SUMMARY OF EXPERIENCE", size=10)
    para = build_summary_third_person(identity, work, full_text, identity.get("position"))
    _add_text(doc, para, size=10)

    # work
    doc.add_paragraph()
    _add_bold_line(doc, "WORK EXPERIENCES", size=10)
    for item in sort_work(work):
        period = _fmt_period(item.get("from")) + " – " + _fmt_period(item.get("to"))
        dur = _dur_ym(item.get("from"), item.get("to"))
        if dur: period = f"{period} — {dur}"
        _add_bold_line(doc, period, size=10)

        role = (item.get("role") or "").strip()
        proj = (item.get("project") or "").strip()
        place = (item.get("city_country") or "").strip()
        line = " — ".join([t for t in [role, f"{proj}, {place}".strip(', ')] if t])
        _add_text(doc, line, size=10)

        raw_bullets = [b for b in (item.get("bullets") or []) if b and str(b).strip()]
        spec_line = project_specs(role, proj, place, raw_bullets)
        if spec_line:
            _add_text(doc, spec_line, size=10)

        bullets = rewrite_project_bullets(raw_bullets)
        for b in bullets:
            _add_bullet(doc, b, size=10)
        doc.add_paragraph()

    # education
    if edu:
        cleaned_edu = [e for e in edu if (isinstance(e, dict) and any([e.get("degree"), e.get("institution"), e.get("city_country"), e.get("year")])) or (isinstance(e, str) and e.strip())]
        if cleaned_edu:
            _add_bold_line(doc, "EDUCATION", size=10)
            for e in cleaned_edu:
                if isinstance(e, dict):
                    deg = e.get("degree") or "-"; inst = e.get("institution") or "-"; cc = e.get("city_country") or "-"; yr = e.get("year") or "-"
                    _add_bullet(doc, f"{deg} — {inst}, {cc} ({yr})", size=10)
                else:
                    _add_bullet(doc, str(e), size=10)
            doc.add_paragraph()

    # skills
    if skills:
        cleaned_sk = [s for s in skills if s and str(s).strip()]
        if cleaned_sk:
            _add_bold_line(doc, "SKILLS", size=10)
            for s in cleaned_sk[:10]:
                _add_bullet(doc, s, size=10)
            doc.add_paragraph()

    # courses
    if courses:
        cleaned_c = [c for c in courses if c and str(c).strip()]
        if cleaned_c:
            _add_bold_line(doc, "COURSES & SEMINARS", size=10)
            for t in cleaned_c[:10]:
                _add_bullet(doc, t, size=10)

    out_path = tempfile.mktemp(suffix='.docx')
    doc.save(out_path)
    with open(out_path, "rb") as fh:
        return fh.read()

# ------------------ RUN BUTTON ------------------
HERE = os.path.dirname(__file__) if "__file__" in globals() else os.getcwd()
TEMPLATE_PATH = os.path.join(HERE, "template", "CURRICULUM VITAE.docx")
if not os.path.exists(TEMPLATE_PATH):
    alt = os.path.join(HERE, "CURRICULUM VITAE.docx")
    if os.path.exists(alt):
        TEMPLATE_PATH = alt

st.button("Generate DOCX CVs", key="gen_btn")
files = st.file_uploader("Upload one or many CVs", type=["pdf","docx"], accept_multiple_files=True, key="cv_files")

if st.session_state.get("gen_btn"):
    if not os.path.exists(TEMPLATE_PATH):
        st.error("Template not found. Place 'CURRICULUM VITAE.docx' inside ./template/ or beside the script.")
    else:
        out_files = []
        if not files:
            st.warning("Upload at least one PDF/DOCX.")
        else:
            # API key + provider
            provider = st.session_state.get("provider_sel","Google Gemini")
            api_key = st.session_state.get("api_key_input") or ""
            model = st.session_state.get("model_pick","gemini-2.5-flash")
            fallback_pos = st.session_state.get("fallback_pos","Tunneling Professional")

            for f in files:
                st.subheader(f"📄 {f.name}")
                raw = extract_text_any(f)
                if not raw:
                    st.error("Could not extract text. Install PyMuPDF/pdfminer.six for PDF and python-docx for DOCX.")
                    st.markdown("---"); continue
                text = strip_pii(canonicalize_text(raw))
                payload = {"desired_position": (fallback_pos or role_from_text(text)), "resume_text": text}

                with st.spinner(f"Calling {provider}…"):
                    try:
                        if provider.startswith("Google"):
                            time.sleep(0.12)  # soft throttle ~8–9 rpm
                            resp = call_gemini_json(api_key, model, payload)
                        else:
                            resp = call_openai_json(api_key, model, payload)
                    except Exception as e:
                        st.error(f"API error: {e}"); st.markdown("---"); continue

                data = sanitize_cv_json(resp or {})
                identity, profile, work, edu, skills, courses = ensure_schema(data, payload["desired_position"], text)

                try:
                    docx_bytes = render_docx_from_template(TEMPLATE_PATH, identity, profile, work, edu, skills, courses, full_text=text)
                except Exception as e:
                    st.error(f"DOCX error: {e}"); continue

                base_no_ext = os.path.splitext(f.name)[0]
                out_name = f"CV BOT - {base_no_ext}.docx"

                st.download_button(
                    label=f"⬇️ Download {out_name}",
                    data=docx_bytes,
                    file_name=out_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_{hashlib.sha1((base_no_ext).encode()).hexdigest()[:8]}"
                )

                tmp_path = os.path.join(tempfile.gettempdir(), out_name)
                with open(tmp_path, "wb") as fh: fh.write(docx_bytes)
                out_files.append(tmp_path)

            if out_files and st.session_state.get("zip_all", True):
                buf = io.BytesIO(); import zipfile
                with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
                    for pth in out_files: z.write(pth, arcname=os.path.basename(pth))
                st.download_button("📦 Download ALL (ZIP)", data=buf.getvalue(), file_name="cv_bot_docx_summaries.zip", mime="application/zip", key="zip_dl")

