
import re
# ====== THIRD-PERSON SUMMARY (HR style, tunnelling-specific; evidence-only) ======
def _years_only(total_months: int, work: list) -> int:
    """Return years as an integer (floor). Fallback by deriving from work durations."""
    try:
        m = int(total_months or 0)
    except Exception:
        m = 0
    if m <= 0 and work:
        # derive crude months from work blocks if needed
        months = 0
        for w in work:
            a, b = w.get("from"), w.get("to")
            da = _parse_date_any(a); db = _parse_date_any(b)
            if da and db:
                months += max(0, (db.year - da.year)*12 + (db.month - da.month))
        m = months
    return max(0, m // 12)

_SECTOR_PATTERNS = [
    (re.compile(r'\b(metro|subway|rail|underground|lrt|mrt|stations?)\b', re.I), "metro/rail"),
    (re.compile(r'\b(road|highway|expressway|tunnel\b(?! boring))\b', re.I), "road"),
    (re.compile(r'\b(sewer|sewage|wastewater|storm|drain|water(?!\s*stop)|pipeline|culvert)\b', re.I), "water/sewer"),
    (re.compile(r'\b(power|cable|utility|utilities|electrical)\b', re.I), "utilities/power"),
    (re.compile(r'\b(airport|runway|airfield)\b', re.I), "airport"),
]

def _infer_sectors(text: str) -> list:
    seen = []
    for rx, label in _SECTOR_PATTERNS:
        if rx.search(text or ""):
            seen.append(label)
    # de-dupe preserving order
    out, seen_set = [], set()
    for s in seen:
        if s not in seen_set:
            out.append(s); seen_set.add(s)
    return out

def _collect_methods_oems_diams_countries(raw_text: str, work: list) -> tuple:
    """
    Collect methods, OEMs, diameters and countries from text/work.
    Falls back to regex-based extraction if the global helpers (_methods/_oems/_diameters_m/_countries)
    are not available in this file.
    """
    blob = raw_text or ""
    for w in work or []:
        parts = [w.get("role",""), w.get("project",""), w.get("city_country","")]
        parts += [b for b in (w.get("bullets") or []) if b]
        blob += " " + " ".join(map(str, parts))

    # METHODS
    try:
        methods = [m.upper() for m in _methods(blob)]
    except Exception:
        # Fallback: use METHOD_HINTS if present
        methods = []
        try:
            for m in METHOD_HINTS:
                rx = r"" + m.replace(" ", "[- ]?") + r""
                if re.search(rx, blob, re.I):
                    methods.append(m.upper())
        except Exception:
            methods = []

    # OEMS
    try:
        oems = _oems(blob)
    except Exception:
        oems = []
        try:
            for o in OEM_HINTS:
                if re.search(r""+re.escape(o)+r"", blob, re.I):
                    oems.append(o)
        except Exception:
            oems = []

    # DIAMETERS
    try:
        diams = _diameters_m(blob)
    except Exception:
        diams = []
        try:
            mm = [float(x)/1000.0 for x in DIAM_MM_RE.findall(blob)]
            m  = [float(x) for x in DIAM_M_RE.findall(blob)]
            diams = sorted(set([round(v,2) for v in mm+m]), reverse=True)
        except Exception:
            diams = []

    # COUNTRIES
    try:
        countries = _countries(blob)
    except Exception:
        countries = []
        try:
            for c in COUNTRY_HINTS:
                if re.search(r""+c+r"", blob, re.I):
                    countries.append(c.replace("\\",""))
            # de-dupe preserve order
            seen=set(); countries=[x for x in countries if not (x in seen or seen.add(x))]
        except Exception:
            countries = []

    return methods, oems, diams, countries

def build_summary_third_person(ident: dict, work: list, raw_text: str, fallback_position: str) -> str:
    """
    Compose a 3–4 sentence paragraph in third person, HR tone.
    Uses only evidenced facts (methods, max diameter, OEMs, sectors, countries).
    Role is taken from fallback_position (UI), with a fallback to identity.position.
    """
    role = (fallback_position or ident.get("position") or "Tunnelling Professional").strip()
    years = _years_only(ident.get("total_experience_months"), work)
    methods, oems, diams, countries = _collect_methods_oems_diams_countries(raw_text, work)
    sectors = _infer_sectors(raw_text + " " + " ".join([str(b) for it in (work or []) for b in (it.get('bullets') or [])]))

    # sentence 1
    s1 = f"Highly experienced {role} with over {years} years in mechanized tunnelling." if years > 0 else f"Highly experienced {role} in mechanized tunnelling."

    # sentence 2 (methods + diameter + OEMs), only if any present
    clauses = []
    if methods:
        clauses.append(", ".join(sorted(set(methods), key=lambda x: methods.index(x))[:3]))
    if diams:
        mx = diams[0]
        try:
            if mx >= 1.0:
                clauses.append(f"up to {mx:.0f} m")
            else:
                clauses.append(f"up to {mx:.2f} m")
        except Exception:
            pass
    s2 = ""
    if clauses:
        s2 = "Demonstrated capability across " + clauses[0]
        if len(clauses) >= 2:
            s2 += f", {', '.join(clauses[1:])}"
        s2 += "."
    if oems:
        s2 = (s2 + " " if s2 else "") + "Experience with OEMs such as " + ", ".join(oems[:3]) + "."

    # sentence 3 (sectors + countries)
    s3 = ""
    if sectors and countries:
        s3 = f"Track record across {', '.join(sectors[:3])} projects in {', '.join(countries[:5])}."
    elif countries:
        s3 = f"International project exposure in {', '.join(countries[:5])}."
    elif sectors:
        s3 = f"Project background spans {', '.join(sectors[:3])}."

    # Merge sentences (3–4 sentences; we have up to 3 deterministic ones)
    parts = [s for s in [s1, s2, s3] if s]
    paragraph = " ".join(parts).strip()

    # guard: force third person (no first-person pronouns)
    if re.search(r'\b(I|my|me|mine)\b', paragraph, re.I):
        paragraph = re.sub(r'\b(I|my|me|mine)\b', 'the candidate', paragraph, flags=re.I)

    return paragraph
# ====== /THIRD-PERSON SUMMARY ======

import os, io, re, json, tempfile, hashlib, unicodedata, datetime, time
from typing import List, Optional, Tuple
import streamlit as st
import time

# ---------------- Ingestion deps ----------------
try:
    import fitz  # PyMuPDF for PDF
    HAVE_PYMUPDF = True
except Exception:
    HAVE_PYMUPDF = False

try:
    from pdfminer_high_level import extract_text as pdfminer_extract_text  # fallback alias if user has different import
    HAVE_PDFMINER = True
except Exception:
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract_text
        HAVE_PDFMINER = True
    except Exception:
        HAVE_PDFMINER = False

try:
    from docx import Document as DocxReader  # python-docx for reading .docx text
    HAVE_DOXC_READ = True
except Exception:
    HAVE_DOXC_READ = False

# ---------------- Word export deps ----------------
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------- Fonts & layout ----------------
COLON_TAB_INCH = 1.60  # tab stop for colon alignment
TOP_FONT = 'Cambria'    # For heading + top identity block
BODY_FONT = 'Verdana'   # For everything else

# ---------------- Regexes & hints ----------------
SPACE_RE = re.compile(r"\s+")
HYPHEN_WRAP_RE = re.compile(r"(\w)-\n(\w)")
NEWLINE_BULLET_RE = re.compile(r"[\u2022\u25CF\u25A0\u00B7]")
HEADER_FOOTER_RE = re.compile(r"Page\s+\d+\s+of\s+\d+", re.I)
EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(r"(\+?\d[\d \-\(\)]{7,}\d)")
URL_RE = re.compile(r"(https?://\S+|www\.\S+)", re.I)

ISO_COUNTRIES = ["Afghanistan", "Albania", "Algeria", "Andorra", "Angola", "Antigua and Barbuda", "Argentina", "Armenia", "Australia", "Austria", "Azerbaijan", "Bahamas", "Bahrain", "Bangladesh", "Barbados", "Belarus", "Belgium", "Belize", "Benin", "Bhutan", "Bolivia", "Bosnia and Herzegovina", "Botswana", "Brazil", "Brunei", "Bulgaria", "Burkina Faso", "Burundi", "Cabo Verde", "Cambodia", "Cameroon", "Canada", "Central African Republic", "Chad", "Chile", "China", "Colombia", "Comoros", "Congo", "Costa Rica", "Côte d’Ivoire", "Croatia", "Cuba", "Cyprus", "Czechia", "Democratic Republic of the Congo", "Denmark", "Djibouti", "Dominica", "Dominican Republic", "Ecuador", "Egypt", "El Salvador", "Equatorial Guinea", "Eritrea", "Estonia", "Eswatini", "Ethiopia", "Fiji", "Finland", "France", "Gabon", "Gambia", "Georgia", "Germany", "Ghana", "Greece", "Grenada", "Guatemala", "Guinea", "Guinea-Bissau", "Guyana", "Haiti", "Honduras", "Hungary", "Iceland", "India", "Indonesia", "Iran", "Iraq", "Ireland", "Israel", "Italy", "Jamaica", "Japan", "Jordan", "Kazakhstan", "Kenya", "Kiribati", "Korea, North", "Korea, South", "Kuwait", "Kyrgyzstan", "Laos", "Latvia", "Lebanon", "Lesotho", "Liberia", "Libya", "Liechtenstein", "Lithuania", "Luxembourg", "Madagascar", "Malawi", "Malaysia", "Maldives", "Mali", "Malta", "Marshall Islands", "Mauritania", "Mauritius", "Mexico", "Micronesia", "Moldova", "Monaco", "Mongolia", "Montenegro", "Morocco", "Mozambique", "Myanmar", "Namibia", "Nauru", "Nepal", "Netherlands", "New Zealand", "Nicaragua", "Niger", "Nigeria", "North Macedonia", "Norway", "Oman", "Pakistan", "Palau", "Panama", "Papua New Guinea", "Paraguay", "Peru", "Philippines", "Poland", "Portugal", "Qatar", "Romania", "Russia", "Rwanda", "Saint Kitts and Nevis", "Saint Lucia", "Saint Vincent and the Grenadines", "Samoa", "San Marino", "Sao Tome and Principe", "Saudi Arabia", "Senegal", "Serbia", "Seychelles", "Sierra Leone", "Singapore", "Slovakia", "Slovenia", "Solomon Islands", "Somalia", "South Africa", "South Sudan", "Spain", "Sri Lanka", "Sudan", "Suriname", "Sweden", "Switzerland", "Syria", "Taiwan", "Tajikistan", "Tanzania", "Thailand", "Timor-Leste", "Togo", "Tonga", "Trinidad and Tobago", "Tunisia", "Turkey", "Turkmenistan", "Tuvalu", "Uganda", "Ukraine", "United Arab Emirates", "United Kingdom", "United States", "Uruguay", "Uzbekistan", "Vanuatu", "Vatican City", "Venezuela", "Vietnam", "Yemen", "Zambia", "Zimbabwe", "Hong Kong", "Macau", "Palestine"]
ALIAS_TO_COUNTRY = {"UK": "United Kingdom", "Great Britain": "United Kingdom", "GB": "United Kingdom", "UAE": "United Arab Emirates", "U.A.E": "United Arab Emirates", "U.A.E.": "United Arab Emirates", "USA": "United States", "U.S.A": "United States", "U.S.A.": "United States", "US": "United States", "U.S.": "United States", "United States of America": "United States", "KSA": "Saudi Arabia", "PRC": "China", "People's Republic of China": "China", "Mainland China": "China", "South Korea": "Korea, South", "Republic of Korea": "Korea, South", "ROK": "Korea, South", "North Korea": "Korea, North", "DPRK": "Korea, North", "Czech Republic": "Czechia", "Ivory Coast": "Côte d’Ivoire", "Cote d'Ivoire": "Côte d’Ivoire", "Burma": "Myanmar", "Lao PDR": "Laos", "Viet Nam": "Vietnam", "Russian Federation": "Russia", "Swaziland": "Eswatini", "Cape Verde": "Cabo Verde", "Taipei, Taiwan": "Taiwan"}
COUNTRY_HINTS = ISO_COUNTRIES  # keep prior inference paths working
LANG_HINTS = ["English","Thai","Hindi","Malay","Tamil","Arabic","Italian","French","German","Spanish","Chinese","Cantonese","Mandarin","Vietnamese","Korean","Japanese"]
OEM_HINTS = ["Herrenknecht","Robbins","Terratec","CREG","Iseki","RASA","Kawasaki","Hitachi Zosen"]

MONTHS = ["JAN","FEB","MAR","APR","MAY","JUN","JUL","AUG","SEP","OCT","NOV","DEC"]
MONTH_MAP = {m:i+1 for i,m in enumerate(MONTHS)}

# ---------------- Utils ----------------
@st.cache_data(show_spinner=False)
def _extract_text_pymupdf(b: bytes) -> str:
    try:
        import fitz as _fitz
        doc = _fitz.open(stream=b, filetype="pdf")
        return "\n".join(p.get_text("text") for p in doc)
    except Exception:
        return ""

@st.cache_data(show_spinner=False)
def _extract_text_pdfminer(b: bytes) -> str:
    try:
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(b); tmp.flush(); path = tmp.name
        try:
            return pdfminer_extract_text(path) or ""
        finally:
            try: os.unlink(path)
            except Exception: pass
    except Exception:
        return ""

def _extract_text_docx(b: bytes) -> str:
    if not HAVE_DOXC_READ: return ""
    from docx import Document as _Doc
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(b); tmp.flush(); path = tmp.name
    try:
        d = _Doc(path)
        texts = []
        for p in d.paragraphs:
            texts.append(p.text)
        for tbl in d.tables:
            for row in tbl.rows:
                texts.append(" | ".join([c.text for c in row.cells]))
        return "\n".join(texts)
    except Exception:
        return ""
    finally:
        try: os.unlink(path)
        except Exception: pass

def extract_text_any(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    data = uploaded_file.getvalue()
    if name.endswith(".pdf"):
        txt = ""
        if HAVE_PYMUPDF:
            txt = _extract_text_pymupdf(data)
        if (not txt or len(txt.strip()) < 100) and HAVE_PDFMINER:
            txt = _extract_text_pdfminer(data)
        return (txt or "").strip()
    elif name.endswith(".docx"):
        return (_extract_text_docx(data) or "").strip()
    else:
        return ""

def canonical_symbols(s: str) -> str:
    s = s.replace("\u2300","Ø").replace("⌀","Ø").replace("Φ","Ø").replace("φ","Ø")
    s = re.sub(r"\b[Dd][Ii][Aa]\.?", " Ø ", s)
    return s

def canonicalize_text(raw: str) -> str:
    s = unicodedata.normalize("NFKC", raw or "")
    s = HEADER_FOOTER_RE.sub(" ", s)
    s = NEWLINE_BULLET_RE.sub(" • ", s)
    s = HYPHEN_WRAP_RE.sub(r"\1\2", s)
    s = canonical_symbols(s)
    s = s.replace("\r","")
    s = SPACE_RE.sub(" ", s)
    return s.strip()

def strip_pii(s: str) -> str:
    s = EMAIL_RE.sub("[REDACTED_EMAIL]", s)
    s = PHONE_RE.sub("[REDACTED_PHONE]", s)
    s = URL_RE.sub("[REDACTED_URL]", s)
    return s

TBM_TERMS = ["tbm operator","tbm pilot","epb","slurry","mixshield","herrenknecht","robbins","terratec","creg","natm","drill & blast","drill and blast"]
def role_from_text(t: str) -> str:
    tl = t.lower()
    if sum(tl.count(k) for k in TBM_TERMS) > 0:
        return "TBM OPERATOR"
    return "Tunneling Professional"


def _fmt_period(s: Optional[str]) -> str:
    """Normalize to 'Month YYYY'. Accept 'YYYY-MM', 'MM/YYYY', 'Mon YYYY', 'Present'."""
    import calendar, re
    if not s or s in ["-","null","None"]:
        return "-"
    s = str(s).strip()
    if re.search(r"present|current|ongoing", s, re.I):
        return "Present"
    m = re.match(r"^(\d{4})[-/](\d{1,2})$", s)  # YYYY-MM
    if m:
        y, mo = int(m.group(1)), int(m.group(2))
        mo = calendar.month_name[mo] if 1 <= mo <= 12 else "--"
        return f"{mo} {y}"
    m = re.match(r"^(\d{1,2})[-/](\d{4})$", s)  # MM/YYYY
    if m:
        mo, y = int(m.group(1)), int(m.group(2))
        mo = calendar.month_name[mo] if 1 <= mo <= 12 else "--"
        return f"{mo} {y}"
    m = re.match(r"^([A-Za-z]{3,9})\s+(\d{4})$", s)  # Mon YYYY or Month YYYY
    if m:
        token = m.group(1)
        try:
            from calendar import month_name, month_abbr
            lookup = {ab.lower():month_name[i] for i,ab in enumerate(month_abbr) if ab}
            lookup.update({mn.lower():mn for mn in month_name if mn})
            return f"{lookup.get(token.lower(), token.title())} {m.group(2)}"
        except Exception:
            return f"{token.title()} {m.group(2)}"
    return s

def _parse_to_key(s: Optional[str]) -> Tuple[int,int,int]:
    if not s: return (0,0,0)
    ss = str(s).strip()
    if re.search(r"present|current|ongoing", ss, re.I): return (9999,12,1)
    m = re.match(r"^(\d{4})[-/](\d{1,2})$", ss)
    if m: return (int(m.group(1)), int(m.group(2)), 0)
    m = re.match(r"^([A-Za-z]{3,9})\s+(\d{4})$", ss)
    if m:
        y=int(m.group(2)); mo=MONTH_MAP.get(m.group(1)[:3].upper(), 12)
        return (y, mo, 0)
    m = re.search(r"(\d{4})", ss)
    if m: return (int(m.group(1)), 12, 0)
    return (0,0,0)

def sort_work(work: List[dict]) -> List[dict]:
    return sorted(work or [], key=lambda x: _parse_to_key(x.get("to")), reverse=True)

def months_to_ym(m) -> str:
    try: m=int(float(m))
    except Exception: return "0m"
    if m<=0: return "0m"
    y=m//12; mo=m%12
    return (f"{y}y " if y else "") + (f"{mo}m" if mo else "")

# ------- Identity helpers -------
def infer_identity(text: str, position_fallback: str) -> dict:
    name_initials = "—"
    m = re.search(r"\b([A-Z])[a-zA-Z]+\s+([A-Z])[a-zA-Z]+", text)
    if m: name_initials = f"{m.group(1)}.{m.group(2)}."
    nat = "—"
    for c in COUNTRY_HINTS:
        if re.search(r"\b"+c+r"\b", text, re.I): nat = c.replace("\\",""); break
    langs_found = [L for L in LANG_HINTS if re.search(r"\b"+L+r"\b", text, re.I)]
    yob = "-"
    m = re.search(r"\b(19|20)\d{2}\b", text)
    if m: yob = m.group(0)
    return {
        "name_initials": name_initials,
        "position": position_fallback,
        "nationality": nat,
        "languages": langs_found or ["English"],
        "year_of_birth": yob,
        "total_experience_months": 0
    }

DIAM_MM_RE = re.compile(r"\b(?:Ø|diam(?:eter)?|dia)\s*([0-9]{3,5})\s*mm\b", re.I)
DIAM_M_RE  = re.compile(r"\b(?:Ø|diam(?:eter)?|dia)\s*([0-9]+(?:\.[0-9]+)?)\s*m\b", re.I)

def _diameters_m(text: str) -> List[float]:
    mm = [float(x)/1000.0 for x in DIAM_MM_RE.findall(text)]
    m  = [float(x) for x in DIAM_M_RE.findall(text)]
    return sorted(set([round(v,2) for v in mm+m]), reverse=True)

def _oems(text: str) -> List[str]:
    found = []
    for o in OEM_HINTS:
        if re.search(r"\b"+re.escape(o)+r"\b", text, re.I):
            found.append(o)
    return sorted(set(found))


def _countries(text: str) -> list:
    # ISO coverage + aliases; case-insensitive whole-word-ish matching.
    t = text or ""
    found = set()
    # Normalize common dotted abbreviations
    norm = t.replace("U.A.E.", "UAE").replace("U.S.A.", "USA")
    # 1) alias matches
    for alias, canonical in ALIAS_TO_COUNTRY.items():
        pat = r'(?<![A-Za-z])' + re.escape(alias) + r'(?![A-Za-z])'
        if re.search(pat, norm, flags=re.IGNORECASE):
            found.add(canonical)
    # 2) direct ISO country names
    for cname in ISO_COUNTRIES:
        pat = r'(?<![A-Za-z])' + re.escape(cname) + r'(?![A-Za-z])'
        if re.search(pat, norm, flags=re.IGNORECASE):
            found.add(cname)
    return sorted(found)


def _parse_date_any(s: Optional[str]) -> Optional[datetime.date]:
    if not s: return None
    s = str(s).strip()
    if re.search(r"present|current|ongoing", s, re.I):
        today = datetime.date.today()
        return datetime.date(today.year, today.month, 1)
    m = re.match(r"^(\d{4})[-/](\d{1,2})$", s)
    if m:
        return datetime.date(int(m.group(1)), int(m.group(2)), 1)
    m = re.match(r"^([A-Za-z]{3,9})\s+(\d{4})$", s)
    if m:
        mo = MONTH_MAP.get(m.group(1)[:3].upper(), 1)
        return datetime.date(int(m.group(2)), mo, 1)
    m = re.search(r"(\d{4})", s)
    if m:
        return datetime.date(int(m.group(1)), 1, 1)
    return None

def _dur_ym(a: Optional[str], b: Optional[str]) -> str:
    da = _parse_date_any(a); db = _parse_date_any(b)
    if not da or not db: return ""
    total = (db.year - da.year)*12 + (db.month - da.month)
    if total < 0: total = 0
    y = total // 12; mo = total % 12
    if y and mo: return f"{y}y {mo}m"
    if y: return f"{y}y"
    if mo: return f"{mo}m"
    return "0m"

# -------- Evidence-gated METHOD detection (project-local) --------

CANON_METHODS = [
    "EPB","SLURRY","MIXSHIELD","NATM","DRILL & BLAST","HARD ROCK","OPEN TBM","SINGLE SHIELD","DOUBLE SHIELD","MICROTUNNELLING","ROADHEADER","RAISE BORING"
]

NAME_SYNONYMS = {
    "EPB": [r"epb\b", r"earth\s*pressure\s*balance", r"epbm\b"],
    "SLURRY": [r"\bslurry\b(?!.*mix)", r"slurry\s*(shield|tbm)", r"bentonite"],
    "MIXSHIELD": [r"mix[-\s]?shield", r"mixshield"],
    "NATM": [r"\bnatm\b", r"\bsem\b", r"sprayed\s*concrete\s*lining", r"\bscl\b"],
    "DRILL & BLAST": [r"\bdrill(?:\s*&\s*| and )blast\b", r"\bd&b\b"],
    "HARD ROCK": [r"hard[-\s]?rock\s*(tbm)?"],
    "OPEN TBM": [r"open\s*(mode|face)\s*tbm", r"\bopen\s*mode\b"],
    "SINGLE SHIELD": [r"single\s*shield"],
    "DOUBLE SHIELD": [r"double\s*shield"],
    "MICROTUNNELLING": [r"\bmicrotunnel", r"\bmtbm\b", r"\bavn\b", r"pipe\s*jacking"],
    "ROADHEADER": [r"\broadheader\b", r"boom\s*header"],
    "RAISE BORING": [r"raise\s*bor(e|ing)"]
}

TRAIT_SYNONYMS = {
    "EPB": [r"face\s*pression|face\s*pressure", r"screw\s*conveyor", r"foam|polymer", r"bulkhead"],
    "SLURRY": [r"separation\s*plant", r"slurry\s*density|viscosity", r"bentonite", r"slurry\s*pump"],
    "MIXSHIELD": [r"air\s*cushion", r"submerged\s*wall", r"compressed\s*air"],
    "NATM": [r"shotcrete", r"lattice\s*girder", r"rock\s*bolt", r"convergence\s*monitor"],
    "DRILL & BLAST": [r"charging", r"initiation", r"stemming", r"blast\s*round"],
    "HARD ROCK": [r"disc\s*cutter", r"penetration\s*rate", r"torque|thrust", r"abrasiv"],
    "OPEN TBM": [r"gripper\s*pads?", r"no\s*pressure\s*control", r"immediate\s*support"],
    "SINGLE SHIELD": [r"shield\s*tail", r"annular\s*grout|tail\s*grout", r"thrust\s*jacks"],
    "DOUBLE SHIELD": [r"telescopic\s*shield", r"simultaneous\s*excavation", r"gripper\s*mode"],
    "MICROTUNNELLING": [r"jacking\s*frame", r"intermediate\s*jacking", r"guidance\s*system"],
    "ROADHEADER": [r"cutter\s*picks?", r"boom", r"profil(?:e|ing)"],
    "RAISE BORING": [r"pilot\s*hole", r"ream(?:er|ing)"]
}

def _regex_any(patterns: List[str]):
    return re.compile("|".join(f"(?:{p})" for p in patterns), re.I)

NAME_REGEX = {k: _regex_any(v) for k,v in NAME_SYNONYMS.items()}
TRAIT_REGEX = {k: _regex_any(v) for k,v in TRAIT_SYNONYMS.items()}

SPECIFICITY = {"MIXSHIELD":3,"SLURRY":2,"EPB":2,"SINGLE SHIELD":2,"DOUBLE SHIELD":2,"OPEN TBM":2,"NATM":2,"DRILL & BLAST":2,"HARD ROCK":2,"MICROTUNNELLING":2,"ROADHEADER":2,"RAISE BORING":2}

def detect_method_project_local(text: str) -> Optional[str]:
    if not text or not text.strip():
        return None
    scores = []
    for label in CANON_METHODS:
        name_hit = bool(NAME_REGEX.get(label, re.compile(r"$^")).search(text))
        trait_hit = bool(TRAIT_REGEX.get(label, re.compile(r"$^")).search(text))
        score = 2 if (name_hit and trait_hit) else (1 if name_hit else 0)
        if score > 0:
            scores.append((score, SPECIFICITY.get(label,1), label))
    if not scores:
        return None
    scores.sort(reverse=True)
    best = scores[0][2]
    if best == "SLURRY":
        if any(lab == "MIXSHIELD" and sc >= scores[0][0] for sc,sp,lab in scores):
            best = "MIXSHIELD"
    return best

# -------- Bullets rewrite (sentence-aware, soft 18-40 words target, 3-6 bullets) --------

ADMIN_NOISE = re.compile(r"\b(email|microsoft (office|windows)|excel|word|ppt|powerpoint|outlook|generic reporting|documentation)\b", re.I)
TUNNEL_SIGNALS = re.compile(r"\b(EPB|slurry|mix[- ]?shield|NATM|drill(?:\s*&\s*blast| and blast)|ring build|VMT|foam|polymer|face pressure|screw conveyor|hyperbaric|cutterhead|convergence|shotcrete|rock bolt|separation plant|slurry density|viscosity|settlement|annular grout|thrust|torque|advance rate|downtime)\b", re.I)

CONJ_SPLIT = re.compile(r"\s+(and|which|that|while|whereas|as well as)\s+", re.I)

def _normalize_whitespace(s: str) -> str:
    return SPACE_RE.sub(" ", (s or "")).strip(" .;,-")

def _word_count(s: str) -> int:
    return 0 if not s else len(_normalize_whitespace(s).split())

def _split_long_sentence(s: str, max_words=42) -> List[str]:
    s = _normalize_whitespace(s)
    if _word_count(s) <= max_words:
        return [s]
    # heuristic split
    tokens = s.split()
    mid = len(tokens)//2
    return [" ".join(tokens[:mid]), " ".join(tokens[mid:])]

def rewrite_project_bullets(raw_bullets: List[str]) -> List[str]:
    # Evidence-only, grammar-preserving bullets. Target 18-40 words; 3-6 bullets.
    cleaned = []
    for b in raw_bullets or []:
        t = _normalize_whitespace(str(b))
        if not t:
            continue
        if ADMIN_NOISE.search(t) and not re.search(r"\b(RAMS|permit|confined|hyperbaric|H2S|CH4|gas|TBM rescue)\b", t, re.I):
            continue
        parts = _split_long_sentence(t, max_words=42)
        for p in parts:
            if _word_count(p) < 8:
                continue
            cleaned.append(p)

    # de-dup (case-insensitive)
    dedup = []
    seen = set()
    for x in cleaned:
        k = _normalize_whitespace(x).lower()
        if k in seen: 
            continue
        seen.add(k)
        dedup.append(x)

    def score(x: str):
        has_sig = 1 if TUNNEL_SIGNALS.search(x) else 0
        has_num = 1 if re.search(r"\d", x) else 0
        wc = _word_count(x)
        length_penalty = abs(wc - 28)  # center near 28 words
        return (-has_sig, -has_num, length_penalty, len(x))

    dedup.sort(key=score)
    return dedup[:6] if len(dedup) >= 6 else dedup

# -------- Summary paragraph (richer: 2-3 sentences) --------



def build_summary_paragraph(full_text: str, work: List[dict], ident: dict) -> str:
    """LLM-written, tunnelling-HR paragraph (3–4+ sentences), using ONE primary role and years-only experience.
       Falls back to deterministic paragraph if the LLM is unavailable or violates rules.
    """
    import re, json

    # --- helper: canonicalise and choose ONE primary role ---
    role_syn = {
        "tbm pilot": "TBM Operator",
        "tbm op": "TBM Operator",
        "tbm operator": "TBM Operator",
        "shift engineer": "TBM Shift Engineer",
        "tbm shift engineer": "TBM Shift Engineer",
        "natm engineer": "NATM Engineer",
        "scl engineer": "NATM Engineer",
        "blasting engineer": "Blasting Engineer",
        "blast engineer": "Blasting Engineer",
        "tunnelling engineer": "Tunnel Engineer",
        "tunnel engineer": "Tunnel Engineer",
        "senior tunnel engineer": "Senior Tunnel Engineer",
        "site engineer": "Tunnel Site Engineer",
        "geotechnical engineer": "Geotechnical Engineer",
        "construction manager": "Construction Manager",
        "tunnel construction manager": "Tunnel Construction Manager",
        "project manager": "Project Manager",
        "project director": "Project Director"
    }
    seniority_order = [
        "Project Director",
        "Construction Manager",
        "Tunnel Construction Manager",
        "Project Manager",
        "Senior Tunnel Engineer",
        "NATM Engineer",
        "TBM Shift Engineer",
        "Geotechnical Engineer",
        "Tunnel Site Engineer",
        "Tunnel Engineer",
        "TBM Operator",
    ]
    rank = {r:i for i,r in enumerate(seniority_order)}

    def norm_role(r):
        rl = (r or "").strip()
        low = rl.lower()
        return role_syn.get(low, rl)

    # parse dates for recency scoring
    def _parse_date_any(sv):
        if not sv: return None
        sv = str(sv).strip()
        if re.search(r"present|current|ongoing", sv, re.I):
            today = datetime.date.today()
            return datetime.date(today.year, today.month, 1)
        m = re.match(r"^(\d{4})[-/](\d{1,2})$", sv)
        if m:
            return datetime.date(int(m.group(1)), int(m.group(2)), 1)
        m = re.match(r"^([A-Za-z]{3,9})\s+(\d{4})$", sv)
        if m:
            MONTH_MAP = {"JAN":1,"FEB":2,"MAR":3,"APR":4,"MAY":5,"JUN":6,"JUL":7,"AUG":8,"SEP":9,"OCT":10,"NOV":11,"DEC":12}
            mo = MONTH_MAP.get(m.group(1)[:3].upper(), 1)
            return datetime.date(int(m.group(2)), mo, 1)
        m = re.search(r"(\d{4})", sv)
        if m:
            return datetime.date(int(m.group(1)), 1, 1)
        return None

    roles = []
    for it in work or []:
        r = norm_role(it.get("role",""))
        if r: roles.append((r, _parse_date_any(it.get("to")), _parse_date_any(it.get("from"))))
    # choose best: seniority -> recency -> frequency
    if roles:
        # frequency map
        freq = {}
        for r,_,_ in roles:
            freq[r] = freq.get(r,0)+1
        # pick max by (seniority rank, recency date, frequency)
        def key_fn(tup):
            r, to, _ = tup
            s_rank = -(100 - rank.get(r, 100))  # lower better; invert for max
            rec = to or datetime.date(1900,1,1)
            return (s_rank, rec, freq.get(r,0))
        best = max(roles, key=key_fn)[0]
        primary_role = best
    else:
        primary_role = ident.get("position") or "Tunnelling Professional"

    # --- collect evidence from work blocks (evidence-only) ---
    DIAM_MM_RE = re.compile(r"\b(?:Ø|diam(?:eter)?|dia)\s*([0-9]{3,5})\s*mm\b", re.I)
    DIAM_M_RE  = re.compile(r"\b(?:Ø|diam(?:eter)?|dia)\s*([0-9]+(?:\.[0-9]+)?)\s*m\b", re.I)
    OEM_HINTS = ["Herrenknecht","Robbins","Terratec","CREG","Iseki","RASA","Kawasaki","Hitachi Zosen"]
    # recognise but do not force—presence only if text contains it
    METHOD_HINTS = [
        "EPB","SLURRY","MIXSHIELD","SINGLE SHIELD","DOUBLE SHIELD",
        "OPEN TBM","HARD ROCK","MICROTUNNELLING","NATM","DRILL & BLAST","SCL"
    ]
    COUNTRY_HINTS = [
        "Qatar","U\\.A\\.E","UAE","United Arab Emirates","Bahrain","Saudi","Oman","Kuwait",
        "Malaysia","Thailand","Singapore","Italy","India","Australia","United Kingdom","UK",
        "Germany","France","Spain","Turkey","Indonesia","Philippines","Vietnam","China",
        "Hong Kong","Taiwan","Japan","Korea","Norway","Poland"
    ]

    def _diameters_m(blob):
        mm = [float(x)/1000.0 for x in DIAM_MM_RE.findall(blob)]
        m  = [float(x) for x in DIAM_M_RE.findall(blob)]
        return sorted(set([round(v,2) for v in (mm+m)]))

    methods, oems, diams, countries = [], [], [], []
    sectors = set()

    for item in work or []:
        role  = str(item.get("role",""))
        proj  = str(item.get("project",""))
        place = str(item.get("city_country",""))
        rb    = [str(b) for b in (item.get("bullets") or []) if b]
        blob  = " ".join([role, proj, place] + rb)

        for mh in METHOD_HINTS:
            if re.search(rf"\b{mh}\b", blob, re.I):
                methods.append(mh)
        oems.extend([o for o in OEM_HINTS if re.search(rf"\b{o}\b", blob, re.I)])
        diams.extend(_diameters_m(blob))
        for c in COUNTRY_HINTS:
            if re.search(rf"\b{c}\b", blob, re.I):
                countries.append(c.replace("\\",""))
        if re.search(r"\bmetro|rail|station\b", blob, re.I): sectors.add("metro/rail")
        if re.search(r"\bwater|sewage|utility\b", blob, re.I): sectors.add("utility/water")
        if re.search(r"\bhighway|road\b", blob, re.I): sectors.add("highway")
        if re.search(r"\bcavern|station box|underground station\b", blob, re.I): sectors.add("caverns/stations")

    # de-dup and trim evidence
    def _uniq(seq): 
        out=[]; 
        [out.append(x) for x in seq if x not in out]; 
        return out

    methods   = _uniq(methods)
    oems      = _uniq(oems)[:3]
    countries = _uniq(countries)[:6]
    diams     = sorted(set(diams))
    years_total = ident.get("total_experience_months") or 0
    try:
        years_total = int(float(years_total))
    except Exception:
        years_total = 0
    years = years_total // 12
    rem = years_total % 12
    years_phrase = f"over {years} years" if rem >= 6 and years>0 else (f"{years} years" if years>0 else "years")

    # --- LLM prompt (tunnelling HR persona) ---
    evidence = {
        "primary_role": primary_role,
        "total_experience_months": years_total,
        "methods": methods,
        "diameters_m": diams,
        "oems": oems,
        "sectors": sorted(list(sectors)),
        "countries": countries,
        "years_phrase_hint": years_phrase
    }
    prompt = (
        "You are a tunnelling HR specialist. Write a single-paragraph 'Summary of Experience' with UK spelling.\n"
        "Length 3–4 sentences by default; go longer only if the evidence is rich.\n"
        "Use only the facts provided; do not invent roles, methods, diameters, OEMs, sectors, or locations.\n"
        "Use the exact primary_role; state years of experience (not months): say 'over X years' if ≥6 extra months, else 'X years'.\n"
        "If present in evidence, mention methods (EPB, slurry, Mixshield, single/double-shield, NATM/SCL, drill-and-blast), "
        "diameters (e.g., 'Ø 6–11 m' or 'up to Ø 11 m'), top OEMs (≤3), sectors/assets (metro/rail, water/utility, highway, station caverns/shafts), and countries/regions.\n"
        f"Evidence JSON: {json.dumps(evidence, ensure_ascii=False)}\n"
        "Output: one paragraph, no bullets, no headers."
    )

    # --- Try LLM if provider+api_key exist; fallback to deterministic paragraph ---
    try:
        if 'provider' in globals() and 'api_key' in globals() and api_key:
            if str(globals().get('provider','')).startswith("OpenAI"):
                from openai import OpenAI
                client = OpenAI(api_key=api_key)
                mdl = globals().get('model') or globals().get('pick') or "gpt-4o-mini"
                resp = client.chat.completions.create(
                    model=mdl,
                    messages=[{"role":"user","content":prompt}],
                    temperature=0.2
                )
                text = (resp.choices[0].message.content or "").strip()
            else:
                import google.generativeai as genai
                genai.configure(api_key=api_key)
                mdl = globals().get('model') or globals().get('pick') or "gemini-2.5-flash"
                gmodel = genai.GenerativeModel(mdl)
                # simple RPM guard
                time.sleep(0.12)  # ~8–9 rpm; user asked cap 10 rpm globally
                resp = gmodel.generate_content(prompt, generation_config={"temperature":0.2})
                text = (resp.text or "").strip()
            # guardrails
            bad = (not text) or ("\n•" in text) or ("\n-" in text) or (len(text.split()) < 25)
            if not bad:
                return text
    except Exception:
        pass

    # fallback: deterministic paragraph similar to previous version
    # (short, evidence-only)
    parts = []
    if years>0:
        parts.append(f"{primary_role} with {years_phrase} in tunnelling.")
    else:
        parts.append(f"{primary_role} in tunnelling.")
    if methods:
        parts.append("Experience with " + ", ".join(m for m in methods[:4]) + ".")
    if diams:
        dmin, dmax = min(diams), max(diams)
        if dmin != dmax:
            parts.append(f"Worked across diameters Ø {dmin:.0f}–{dmax:.0f} m.")
        else:
            parts.append(f"Worked up to Ø {dmax:.0f} m.")
    if oems:
        parts.append("OEMs: " + ", ".join(oems) + ".")
    if countries:
        parts.append("Countries: " + ", ".join(countries) + ".")
    if sectors:
        parts.append("Sectors: " + ", ".join(sorted(sectors)) + ".")
    return " ".join(parts)

def _clear_body(doc: Document):
    for para in list(doc.paragraphs):
        p = para._element; p.getparent().remove(p)
    for tbl in list(doc.tables):
        t = tbl._element; t.getparent().remove(t)

def _add_heading(doc: Document, text: str, size=16, bold=True, italic=True, align_center=True):
    p=doc.add_paragraph()
    r=p.add_run(str(text))
    r.font.name = TOP_FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if align_center: p.alignment = 1
    return p

def _add_bold_line(doc: Document, text, size=10):
    p=doc.add_paragraph()
    r=p.add_run(str(text))
    r.bold=True
    r.font.name = BODY_FONT
    r.font.size = Pt(size)

def _add_text(doc: Document, text, size=10):
    p=doc.add_paragraph()
    r=p.add_run(str(text))
    r.font.name = BODY_FONT
    r.font.size = Pt(size)

def _add_bullet(doc: Document, text, size=10):
    try:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(str(text))
    except Exception:
        p = doc.add_paragraph()
        try: 
            p.paragraph_format.left_indent = Inches(0.25)
        except Exception: 
            pass
        r = p.add_run("• " + str(text))
    r.font.name = BODY_FONT
    r.font.size = Pt(size)

def _add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    p_par = p._element
    p_pr = p_par.get_or_add_pPr()
    p_borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    p_borders.append(bottom)
    p_pr.append(p_borders)

def _add_identity_line(doc: Document, label: str, value: str, tab_pos_in=COLON_TAB_INCH):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Pt(0); pf.first_line_indent = Pt(0)
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    pf.line_spacing = 1.0

    pPr = p._p.get_or_add_pPr()
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        tabs = OxmlElement('w:tabs'); pPr.append(tabs)
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'left')
    tab.set(qn('w:pos'), str(int(tab_pos_in * 1440)))
    tab.set(qn('w:leader'), 'underscore')
    tabs.append(tab)

    r1 = p.add_run(label + " ")
    r1.font.name = TOP_FONT; r1.font.size = Pt(10); r1.bold = True

    rtab = p.add_run('\t')
    rtab.font.name = TOP_FONT; rtab.font.size = Pt(10); rtab.bold = True

    r2 = p.add_run(': ')
    r2.font.name = TOP_FONT; r2.font.size = Pt(10); r2.bold = True

    r3 = p.add_run(str(value))
    r3.font.name = TOP_FONT; r3.font.size = Pt(10); r3.bold = True

def _add_top_identity_paragraphs(doc: Document, pos: str, name_i: str, nat: str, langs: str, yob: str, exp: str):
    labels = ["POSITION","NAME","NATIONALITY","LANGUAGES","YEAR OF BIRTH","EXPERIENCE"]
    values = [pos, name_i, nat, langs, yob, exp]
    for L, V in zip(labels, values):
        _add_identity_line(doc, L, V, tab_pos_in=COLON_TAB_INCH)
        doc.add_paragraph("")  # blank line under each label line

# -------- LLM plumbing --------

def _strong_prompt() -> str:
    return (
        "You are a specialized HR in the tunneling industry. "
        "Return STRICT JSON with keys exactly:\n"
        "identity: { name_initials, position, nationality, languages[], year_of_birth, total_experience_months }\n"
        "profile_summary: short paragraph (2-4 lines) summarizing seniority, key methods (EPB/Slurry/Mixshield/NATM/Hard Rock/Open TBM/Shield/Drill & Blast), diameters, TBM OEMs, and countries.\n"
        "work_experiences: array of { from:'YYYY-MM'|Mon YYYY|'-', to:'YYYY-MM'|Mon YYYY|'Present', role, project, city_country, bullets[] }\n"
        "education: string[] or objects with degree/institution/city_country/year\n"
        "skills: string[]\n"
        "courses: string[]\n"
        "Constraints: redact PII; initials for names; use '-' for unknown; do not include extra text outside JSON."
    )

def call_gemini_json(api_key: str, model: str, payload: dict) -> dict:
    import google.generativeai as genai, json as _json, re as _re
    genai.configure(api_key=api_key)
    gmodel = genai.GenerativeModel(model)
    resp = gmodel.generate_content(
        _strong_prompt() + "\nINPUT:\n" + json.dumps(payload, ensure_ascii=False),
        generation_config={"temperature": 0, "response_mime_type":"application/json"}
    )
    out = resp.text or "{}"
    try:
        return _json.loads(out)
    except Exception:
        m = _re.search(r"(\{.*\})", out, _re.S)
        return _json.loads(m.group(1)) if m else {}

def call_openai_json(api_key: str, model: str, payload: dict) -> dict:
    from openai import OpenAI
    client = OpenAI(api_key=api_key)
    resp = client.chat.completions.create(
        model=model,
        messages=[{"role":"system","content":_strong_prompt()},
                  {"role":"user","content": json.dumps(payload, ensure_ascii=False)}],
        response_format={"type":"json_object"},
        temperature=0
    )
    return json.loads(resp.choices[0].message.content or "{}")

# ------------------ APP UI ------------------
st.set_page_config(page_title="CV Summary → DOCX Template", layout="wide")
st.title("CV Summary Maker")

st.sidebar.header("Provider")
provider = st.sidebar.selectbox("Choose API", ["Google Gemini", "OpenAI (ChatGPT)"], index=0, key="provider_sel")
prefill = os.getenv("GEMINI_API_KEY") if provider.startswith("Google") else os.getenv("OPENAI_API_KEY")
api_key = st.sidebar.text_input("API Key", type="password", value=prefill or "", key="api_key_input")

if provider.startswith("Google"):
    models = ["gemini-2.5-flash", "gemini-2.5-pro", "gemini-2.0-flash-exp", "Custom..."]
    pick = st.sidebar.selectbox("Model", models, index=0, key="model_pick")
    model = st.sidebar.text_input("Custom model name", value="", key="custom_model_name") if pick == "Custom..." else pick
else:
    models = ["gpt-4o-mini", "gpt-4o", "gpt-4.1-mini", "Custom..."]
    pick = st.sidebar.selectbox("Model", models, index=0, key="model_pick")
    model = st.sidebar.text_input("Custom model name", value="", key="custom_model_name") if pick == "Custom..." else pick

st.sidebar.header("Options")
default_position = st.sidebar.text_input("Fallback POSITION", value="Tunneling Professional", key="fallback_pos")
batch_zip = st.sidebar.checkbox("Also create ZIP of all DOCXs", value=True, key="zip_all")

HERE = os.path.dirname(__file__) if "__file__" in globals() else os.getcwd()
TEMPLATE_PATH = os.path.join(HERE, "template", "CURRICULUM VITAE.docx")
if not os.path.exists(TEMPLATE_PATH):
    alt = os.path.join(HERE, "CURRICULUM VITAE.docx")
    if os.path.exists(alt):
        TEMPLATE_PATH = alt

st.header("Upload resumes (PDF or DOCX)")
files = st.file_uploader("Upload one or many CVs", type=["pdf","docx"], accept_multiple_files=True, key="cv_files")

def build_payload(text: str, fallback_position: str) -> dict:
    # UI fallback position is the single source of truth; only if it's empty do we infer.
    role = (fallback_position or "").strip()
    if not role:
        role = role_from_text(text)
    return {"desired_position": role, "resume_text": text}

def sanitize_cv_json(data: dict) -> dict:
    s = json.dumps(data, ensure_ascii=False)
    s = EMAIL_RE.sub("[REDACTED_EMAIL]", s)
    s = PHONE_RE.sub("[REDACTED_PHONE]", s)
    s = URL_RE.sub("[REDACTED_URL]", s)
    try: return json.loads(s)
    except Exception: return data

def ensure_schema(d: dict, position: str, raw_text: str):
    d = d or {}
    ident = d.get("identity") or {}
    if not ident or (ident.get("name_initials") in [None,"-","—",""] and ident.get("nationality") in [None,"-","—",""]):
        ident = {**infer_identity(raw_text, position), **ident}
    months = ident.get("total_experience_months")
    try: months = int(float(months))
    except Exception: months = 12 * len(d.get("work_experiences") or [])
    ident["total_experience_months"] = months
    langs = ident.get("languages") or ["English"]
    if isinstance(langs, list) and not langs:
        ident["languages"] = ["English"]
    if isinstance(langs, str) and not langs.strip():
        ident["languages"] = ["English"]
    profile = d.get("profile_summary") or ""
    # Force UI fallback position as single source of truth
    if position and str(position).strip():
        ident["position"] = str(position).strip()
    work = d.get("work_experiences") or []
    edu = d.get("education") or []
    skills = d.get("skills") or []
    courses = d.get("courses") or []
    return ident, profile, work, edu, skills, courses

def project_specs(role, proj, place, bullets) -> str:
    blob = " ".join([str(role or ""), str(proj or ""), str(place or "")] + [str(b) for b in (bullets or []) if b])
    meth = detect_method_project_local(blob)
    ds = _diameters_m(blob)
    os_ = _oems(blob)
    parts = []
    if meth: parts.append("Method: " + meth)
    if ds: parts.append("Ø: " + ", ".join(f"{d:.2f} m" for d in ds[:3]))
    if os_: parts.append("OEM: " + ", ".join(os_[:3]))
    return " | ".join(parts)


# ---------- Summary helpers (synonym-aware; summary-only) ----------
METHOD_FAMILIES = {
    "EPB": r"\b(EPB|earth\s*pressure\s*balance|balance\s*shield)\b",
    "Slurry": r"\b(slurry|bentonite\s*slurry)\b",
    "Mixshield": r"\b(mix\s*shield|mixshield)\b",
    "Double Shield TBM": r"\b(double\s*shield|DS\b|DSTS)\b",
    "Single Shield TBM": r"\b(single\s*shield|SS\b)\b",
    "Open TBM": r"\b(open\s*(?:tbm|gripper|tunnel\s*borer)|gripper\s*tbm)\b",
    "Hard Rock": r"\b(hard\s*rock)\b",
    "NATM/SEM": r"\b(NATM|SEM|drill\s*(?:&|and)\s*blast|D&B)\b",
}

SECTOR_PATTERNS = [
    (re.compile(r"\b(metro|subway|rail|underground|lrt|mrt|stations?)\b", re.I), "metro/rail"),
    (re.compile(r"\b(road|highway|expressway)\b", re.I), "road"),
    (re.compile(r"\b(water|sewer|wastewater|drainage)\b", re.I), "water/sewer"),
    (re.compile(r"\b(airport|runway|terminal)\b", re.I), "airport"),
    (re.compile(r"\b(mining|mine)\b", re.I), "mining"),
    (re.compile(r"\b(hydro|power\s*plant)\b", re.I), "power/hydro"),
    (re.compile(r"\b(utility|cable|pipe\s*jacking)\b", re.I), "utilities"),
]

def _extract_methods_synonyms(text: str) -> list:
    buf = text or ""
    scores = {}
    for label, rx in METHOD_FAMILIES.items():
        n = len(re.findall(rx, buf, flags=re.I))
        if n:
            scores[label] = scores.get(label, 0) + n
    # sort by frequency desc, keep top 3
    ordered = [k for k, _ in sorted(scores.items(), key=lambda kv: (-kv[1], kv[0].lower()))]
    return ordered[:3]

def _extract_sectors(text: str) -> list:
    buf = text or ""
    found = []
    for rx, label in SECTOR_PATTERNS:
        if rx.search(buf):
            found.append(label)
    # preserve order, dedupe
    seen = set(); out = []
    for x in found:
        if x not in seen:
            out.append(x); seen.add(x)
    return out[:3]

def _collect_summary_facts(identity: dict, work: list, raw_text: str) -> dict:
    # Combine text from work items and full text
    blob_parts = [raw_text or ""]
    for w in work or []:
        blob_parts.extend([str(w.get("role","")), str(w.get("project","")), str(w.get("city_country",""))])
        for b in (w.get("bullets") or []):
            if b: blob_parts.append(str(b))
    blob = " ".join(blob_parts)

    # Methods: use synonym-aware first, fallback to existing helpers if present
    methods = _extract_methods_synonyms(blob)
    if not methods:
        try:
            methods = [m.upper() for m in _methods(blob)]
        except Exception:
            methods = []

    # OEMs
    try:
        oems = _oems(blob)
    except Exception:
        oems = []
        for o in OEM_HINTS:
            if re.search(r"\b"+re.escape(o)+r"\b", blob, re.I):
                oems.append(o)
    # cap
    oems = list(dict.fromkeys(oems))[:3]

    # Sectors
    sectors = _extract_sectors(blob)

    # Countries
    try:
        countries = _countries(blob)
    except Exception:
        countries = []
        for c in COUNTRY_HINTS:
            if re.search(r"\b"+c+r"\b", blob, re.I):
                countries.append(c.replace("\\\\",""))
    countries = list(dict.fromkeys(countries))[:5]

    # Years (not months)
    months = int(identity.get("total_experience_months") or 0)
    years = months // 12 if months > 0 else 0

    return {"methods": methods, "oems": oems, "sectors": sectors, "countries": countries, "years": years}

def build_summary_third_person(identity: dict, work: list, raw_text: str, fallback_position: str) -> str:
    # Resolve position: prefer forced identity.position (caller ensures override)
    position = (identity.get("position") or fallback_position or "Tunneling Professional").strip()

    facts = _collect_summary_facts(identity, work, raw_text)
    years = facts["years"]
    methods = facts["methods"]
    oems = facts["oems"]
    sectors = facts["sectors"]
    countries = facts["countries"]

    parts = []

    # Sentence 1
    if years > 0:
        parts.append(f"Highly experienced {position} with over {years} years in mechanized tunnelling.")
    else:
        parts.append(f"Highly experienced {position} in mechanized tunnelling.")

    # Sentence 2 (methods + OEMs) — methods are MUST-HAVE if any evidence
    if methods:
        if oems:
            parts.append(f"Proficient in {', '.join(methods)} and experienced with OEMs such as {', '.join(oems)}.")
        else:
            parts.append(f"Proficient in {', '.join(methods)}.")
    elif oems:
        parts.append(f"Experienced with OEMs such as {', '.join(oems)}.")

    # Sentence 3 (sectors)
    if sectors:
        parts.append(f"Track record across {', '.join(sectors)} projects.")

    # Sentence 4 (countries)
    if countries:
        parts.append(f"International experience in {', '.join(countries)}.")

    return " ".join(parts)


def render_docx_from_template(template_path: str, identity: dict, profile: str, work: List[dict], edu: List, skills: List[str], courses: List[str], full_text: str) -> bytes:
    doc = Document(template_path)
    _clear_body(doc)

    p = doc.add_paragraph()
    r=p.add_run("CURRICULUM VITAE")
    r.font.name = TOP_FONT; r.font.size = Pt(16); r.bold = True; r.italic = True
    p.alignment = 1
    doc.add_paragraph()

    pos = identity.get("position") or "Tunneling Professional"
    name_i = identity.get("name_initials") or "—"
    nat = identity.get("nationality") or "—"
    langs_field = identity.get("languages") or ["English"]
    if isinstance(langs_field, str): 
        langs = langs_field or "English"
    else: 
        langs = ", ".join([str(x) for x in langs_field]) or "English"
    yob = identity.get("year_of_birth") or "—"
    exp = months_to_ym(identity.get("total_experience_months") or 0)

    _add_top_identity_paragraphs(doc, pos, name_i, nat, langs, yob, exp)

    _add_horizontal_rule(doc)
    doc.add_paragraph()

    _add_bold_line(doc, "SUMMARY OF EXPERIENCE", size=10)
    para = build_summary_third_person(identity, work, "", identity.get("position"))
    _add_text(doc, para, size=10)
    doc.add_paragraph()
    _add_bold_line(doc, "WORK EXPERIENCES", size=10)
    for item in sort_work(work):
        period = _fmt_period(item.get("from")) + " – " + _fmt_period(item.get("to"))
        dur = _dur_ym(item.get("from"), item.get("to"))
        if dur: period = f"{period} — {dur}"
        _add_bold_line(doc, period, size=10)

        role = (item.get("role") or "").strip()
        proj = (item.get("project") or "").strip()
        place = (item.get("city_country") or "").strip()
        line = " — ".join([t for t in [role, f"{proj}, {place}".strip(', ')] if t])
        _add_text(doc, line, size=10)

        raw_bullets = [b for b in (item.get("bullets") or []) if b and str(b).strip()]
        spec_line = project_specs(role, proj, place, raw_bullets)
        if spec_line:
            _add_text(doc, spec_line, size=10)

        bullets = rewrite_project_bullets(raw_bullets)
        for b in bullets:
            _add_bullet(doc, b, size=10)
        doc.add_paragraph()

    if edu:
        cleaned_edu = [e for e in edu if (isinstance(e, dict) and any([e.get("degree"), e.get("institution"), e.get("city_country"), e.get("year")])) or (isinstance(e, str) and e.strip())]
        if cleaned_edu:
            _add_bold_line(doc, "EDUCATION", size=10)
            for e in cleaned_edu:
                if isinstance(e, dict):
                    deg = e.get("degree") or "-"
                    inst = e.get("institution") or "-"
                    cc = e.get("city_country") or "-"
                    yr = e.get("year") or "-"
                    _add_bullet(doc, f"{deg} — {inst}, {cc} ({yr})", size=10)
                else:
                    _add_bullet(doc, str(e), size=10)
            doc.add_paragraph()

    if skills:
        cleaned_sk = [s for s in skills if s and str(s).strip()]
        if cleaned_sk:
            _add_bold_line(doc, "SKILLS", size=10)
            for s in cleaned_sk[:10]:
                _add_bullet(doc, s, size=10)
            doc.add_paragraph()

    if courses:
        cleaned_c = [c for c in courses if c and str(c).strip()]
        if cleaned_c:
            _add_bold_line(doc, "COURSES & SEMINARS", size=10)
            for t in cleaned_c[:10]:
                _add_bullet(doc, t, size=10)

    out_path = tempfile.mktemp(suffix='.docx')
    doc.save(out_path)
    with open(out_path, "rb") as fh:
        return fh.read()

HERE = os.path.dirname(__file__) if "__file__" in globals() else os.getcwd()
TEMPLATE_PATH = os.path.join(HERE, "template", "CURRICULUM VITAE.docx")
if not os.path.exists(TEMPLATE_PATH):
    alt = os.path.join(HERE, "CURRICULUM VITAE.docx")
    if os.path.exists(alt):
        TEMPLATE_PATH = alt
# ---- Merge files from multiple uploaders (if present) ----
try:
    _files_main = st.session_state.get("cv_files")
    _files_extra = st.session_state.get("cv_files_extra")
    _merged = []
    if _files_main: 
        _merged.extend(_files_main)
    if _files_extra: 
        _merged.extend(_files_extra)
    if _merged:
        files = _merged
except Exception:
    pass
# ---- /merge ----


if st.button("Generate DOCX CVs", key="gen_btn"):
    if not os.path.exists(TEMPLATE_PATH):
        st.error("Template not found. Place 'CURRICULUM VITAE.docx' inside ./template/ or beside the script.")
    elif not api_key:
        st.error("Paste your API key.")
    elif not files:
        st.warning("Upload at least one PDF/DOCX.")
    else:
        out_files = []
        for f in files:
            st.subheader(f"📄 {f.name}")
            raw = extract_text_any(f)
            if not raw:
                st.error("Could not extract text. Install PyMuPDF/pdfminer.six for PDF and python-docx for DOCX.")
                st.markdown("---"); continue
            text = strip_pii(canonicalize_text(raw))
            payload = build_payload(text, st.session_state.get("fallback_pos","Tunneling Professional"))

            with st.spinner(f"Calling {st.session_state.get('provider_sel','Google Gemini')}…"):
                try:
                    if provider.startswith("Google"):
                        time.sleep(0.15)  # soft throttle
                        resp = call_gemini_json(api_key, model or st.session_state.get("model_pick","gemini-2.5-flash"), payload)
                    else:
                        resp = call_openai_json(api_key, model or st.session_state.get("model_pick","gpt-4o-mini"), payload)
                except Exception as e:
                    st.error(f"API error: {e}"); st.markdown("---"); continue

            data = sanitize_cv_json(resp or {})
            identity, profile, work, edu, skills, courses = ensure_schema(data, payload["desired_position"], text)

            try:
                docx_bytes = render_docx_from_template(TEMPLATE_PATH, identity, profile, work, edu, skills, courses, full_text=text)
            except Exception as e:
                st.error(f"DOCX error: {e}"); continue

            base_no_ext = os.path.splitext(f.name)[0]
            out_name = f"CV BOT - {base_no_ext}.docx"

            st.download_button(
                label=f"⬇️ Download {out_name}",
                data=docx_bytes,
                file_name=out_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_{hashlib.sha1((base_no_ext).encode()).hexdigest()[:8]}"
            )

            tmp_path = os.path.join(tempfile.gettempdir(), out_name)
            with open(tmp_path, "wb") as fh: fh.write(docx_bytes)
            out_files.append(tmp_path)

        if out_files and st.session_state.get("zip_all", True):
            buf = io.BytesIO(); import zipfile
            with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
                for pth in out_files: z.write(pth, arcname=os.path.basename(pth))
            st.download_button("📦 Download ALL (ZIP)", data=buf.getvalue(), file_name="cv_bot_docx_summaries.zip", mime="application/zip", key="zip_dl")
